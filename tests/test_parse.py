"""Unit-тесты на чистые helpers proverka — без сети, без LLM."""
from __future__ import annotations

import io
import sys
from pathlib import Path

# Чтобы импортировать agent.py из родительской папки без установки.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import zipfile

import pytest
from docx import Document

import agent as proverka  # noqa: E402


def test_parse_json_strips_think_block() -> None:
    raw = (
        "<think>модель размышляет...</think>\n"
        '{"overall": 75, "by_criterion": {"topic_match": 80}}'
    )
    assert proverka._parse_json_blob(raw) == {
        "overall": 75,
        "by_criterion": {"topic_match": 80},
    }


def test_parse_json_strips_code_fence() -> None:
    raw = (
        "Вот ответ:\n"
        "```json\n"
        '{"overall": 60, "strengths": ["a", "b"]}\n'
        "```\n"
    )
    assert proverka._parse_json_blob(raw)["overall"] == 60


def test_parse_json_handles_plain_json() -> None:
    raw = '{"recommendation": "доработать", "overall": 50}'
    assert proverka._parse_json_blob(raw)["recommendation"] == "доработать"


def test_extract_text_docx(tmp_path: Path) -> None:
    p = tmp_path / "w.docx"
    d = Document()
    d.add_paragraph("Первый абзац")
    d.add_paragraph("Второй абзац")
    d.save(p)
    text = proverka._extract_text(p)
    assert "Первый абзац" in text
    assert "Второй абзац" in text


def test_extract_text_unsupported_format(tmp_path: Path) -> None:
    p = tmp_path / "x.txt"
    p.write_text("hello")
    assert "неподдерживаемый формат" in proverka._extract_text(p)


def test_list_work_files_flat_layout(tmp_path: Path) -> None:
    (tmp_path / "Иванов.docx").write_bytes(b"")
    (tmp_path / "Петров.pdf").write_bytes(b"")
    items = proverka._list_work_files(tmp_path)
    names = sorted(n for n, _ in items)
    assert names == ["Иванов", "Петров"]


def test_list_work_files_nested_layout(tmp_path: Path) -> None:
    (tmp_path / "Сидоров").mkdir()
    (tmp_path / "Сидоров" / "doc.docx").write_bytes(b"")
    (tmp_path / "Кузнецов").mkdir()
    (tmp_path / "Кузнецов" / "work.pdf").write_bytes(b"")
    items = proverka._list_work_files(tmp_path)
    names = sorted(n for n, _ in items)
    assert names == ["Кузнецов", "Сидоров"]


def test_list_work_files_skips_empty_subdirs(tmp_path: Path) -> None:
    (tmp_path / "Empty").mkdir()
    (tmp_path / "Ivanov.pdf").write_bytes(b"")
    items = proverka._list_work_files(tmp_path)
    assert [n for n, _ in items] == ["Ivanov"]


def test_per_work_doc_renders_all_fields() -> None:
    score = proverka.WorkScore(
        name="work1",
        overall=77,
        by_criterion={
            "topic_match": 90,
            "novelty": 60,
            "method": 70,
            "results": 75,
            "formatting": 85,
        },
        strengths=["Чёткая структура"],
        weaknesses=["Нет stat-анализа"],
        recommendation="Рекомендовать с доработкой",
        raw_excerpt="...",
    )
    blob = proverka._per_work_doc(score, "ВНТК-2025")
    d = Document(io.BytesIO(blob))
    text = "\n".join(p.text for p in d.paragraphs)
    assert "ВНТК-2025" in text
    assert "77 из 100" in text
    assert "Чёткая структура" in text
    assert "Нет stat-анализа" in text
    assert "Рекомендовать с доработкой" in text
