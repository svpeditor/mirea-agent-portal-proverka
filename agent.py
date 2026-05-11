"""proverka — реальный агент проверки конкурсных работ.

Контракт:
- $INPUT_DIR/works/<подпапка-или-файл> — работы школьников (PDF/DOCX)
- params.competition: название конкурса
- params.grade_level: 5-7 / 8-9 / 10-11

Вывод:
- report.docx — сводное заключение с таблицей оценок
- per_work.zip — папка с заключениями по каждой работе

Использует DeepSeek-R1 через OpenRouter (OPENROUTER_API_KEY ephemeral
из env, инжектится порталом для агентов с runtime.llm).
"""
from __future__ import annotations

import io
import json
import os
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

import httpx
import pdfplumber
from docx import Document

from portal_sdk import Agent

CHECKLIST = [
    ("topic_match", "Соответствие теме"),
    ("novelty", "Научная новизна"),
    ("method", "Качество методологии и эксперимента"),
    ("results", "Чёткость и обоснованность результатов"),
    ("formatting", "Оформление и грамотность"),
]

MAX_TEXT_CHARS = 30_000  # один запрос к LLM, длинные работы режем


@dataclass
class WorkScore:
    name: str
    overall: int
    by_criterion: dict[str, int]
    strengths: list[str]
    weaknesses: list[str]
    recommendation: str
    raw_excerpt: str


def _extract_text(path: Path) -> str:
    """Из PDF/DOCX вытащить чистый текст."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        chunks: list[str] = []
        try:
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t:
                        chunks.append(t)
        except Exception as e:  # noqa: BLE001
            return f"[ошибка чтения PDF: {e}]"
        return "\n".join(chunks)
    if ext == ".docx":
        try:
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:  # noqa: BLE001
            return f"[ошибка чтения DOCX: {e}]"
    return f"[неподдерживаемый формат: {ext}]"


def _list_work_files(works_dir: Path) -> list[tuple[str, list[Path]]]:
    """Вернуть список (имя_работы, файлы). Поддерживаем 2 раскладки:
    - works/<автор>/<файл> — подпапка = работа
    - works/<файл> — файл сам по себе = работа
    """
    items: list[tuple[str, list[Path]]] = []
    for entry in sorted(works_dir.iterdir()):
        if entry.is_dir():
            files = sorted(
                p for p in entry.rglob("*")
                if p.is_file() and p.suffix.lower() in (".pdf", ".docx")
            )
            if files:
                items.append((entry.name, files))
        elif entry.is_file() and entry.suffix.lower() in (".pdf", ".docx"):
            items.append((entry.stem, [entry]))
    return items


def _llm_review(text: str, competition: str, grade_level: str, model: str, api_key: str, base_url: str) -> dict:
    """Один запрос к DeepSeek-R1 на разбор работы. Возвращает dict с оценками."""
    criteria_list = "\n".join(f"- {k}: {label}" for k, label in CHECKLIST)
    system = (
        "Ты — научный эксперт, оцениваешь работу школьника на конкурсе. "
        "Отвечай строго в JSON по схеме, без markdown, без лишних слов."
    )
    user = f"""Конкурс: {competition}
Класс участников: {grade_level}

Чек-лист критериев (каждый по шкале 0-100):
{criteria_list}

Текст работы (может быть обрезан):
---
{text[:MAX_TEXT_CHARS]}
---

Верни строго JSON:
{{
  "by_criterion": {{ "<criterion_key>": <0..100>, ... }},
  "overall": <0..100>,
  "strengths": ["...", "..."],
  "weaknesses": ["...", "..."],
  "recommendation": "<1-2 предложения: рекомендовать / доработать / отклонить и почему>"
}}
"""
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.2,
        "max_tokens": 2000,
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://github.com/svpeditor/mirea-agent-portal-proverka",
        "X-Title": "mirea-proverka",
    }
    with httpx.Client(timeout=120) as client:
        r = client.post(f"{base_url.rstrip('/')}/chat/completions", json=payload, headers=headers)
        r.raise_for_status()
        data = r.json()
    content = data["choices"][0]["message"]["content"]
    return _parse_json_blob(content)


def _parse_json_blob(s: str) -> dict:
    """LLM может обернуть JSON в ```json ... ``` или добавить thinking-блок."""
    s = s.strip()
    # убрать <think>...</think> от reasoning-моделей
    s = re.sub(r"<think>.*?</think>", "", s, flags=re.DOTALL).strip()
    # ```json блоки
    m = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", s, re.DOTALL)
    if m:
        s = m.group(1)
    # вынуть первый верхнеуровневый JSON
    m = re.search(r"\{.*\}", s, re.DOTALL)
    if m:
        s = m.group(0)
    return json.loads(s)


def _per_work_doc(score: WorkScore, competition: str) -> bytes:
    doc = Document()
    doc.add_heading(f"Заключение: {score.name}", level=0)
    doc.add_paragraph(f"Конкурс: {competition}")
    doc.add_heading("Итоговая оценка", level=1)
    doc.add_paragraph(f"{score.overall} из 100")

    doc.add_heading("По критериям", level=1)
    t = doc.add_table(rows=1, cols=2)
    h = t.rows[0].cells
    h[0].text = "Критерий"
    h[1].text = "Балл"
    for key, label in CHECKLIST:
        row = t.add_row().cells
        row[0].text = label
        row[1].text = str(score.by_criterion.get(key, "—"))

    if score.strengths:
        doc.add_heading("Сильные стороны", level=1)
        for s in score.strengths:
            doc.add_paragraph(s, style="List Bullet")
    if score.weaknesses:
        doc.add_heading("Слабые места", level=1)
        for w in score.weaknesses:
            doc.add_paragraph(w, style="List Bullet")

    doc.add_heading("Рекомендация эксперта", level=1)
    doc.add_paragraph(score.recommendation or "—")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main() -> None:
    agent = Agent()
    params = agent.params
    competition: str = params.get("competition", "(без названия)")
    grade_level: str = params.get("grade_level", "10-11")

    api_key = os.environ.get("OPENROUTER_API_KEY", "").strip()
    base_url = os.environ.get("OPENROUTER_BASE_URL", "https://openrouter.ai/api/v1").strip()
    model = os.environ.get("LLM_MODEL", "deepseek/deepseek-r1").strip()

    if not api_key:
        agent.failed("OPENROUTER_API_KEY не передан в контейнер — портал должен инжектить ephemeral-токен.")
        return

    agent.log("info", f"proverka: конкурс={competition!r}, класс={grade_level!r}, модель={model}")

    works_dir = agent.input_dir("works")
    items = _list_work_files(works_dir)
    n = len(items)
    if n == 0:
        agent.failed("В папке works не найдено ни одной работы (.pdf/.docx).")
        return

    agent.log("info", f"К проверке принято {n} работ")
    scores: list[WorkScore] = []
    zip_buf = io.BytesIO()

    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, (name, files) in enumerate(items):
            agent.progress(i / n, f"Работа {i + 1}/{n}: {name}")
            text = "\n\n".join(_extract_text(p) for p in files).strip()
            if not text:
                agent.log("warn", f"{name}: пустой текст после извлечения, оценить нельзя")
                score = WorkScore(
                    name=name, overall=0,
                    by_criterion={k: 0 for k, _ in CHECKLIST},
                    strengths=[],
                    weaknesses=["Не удалось извлечь текст из файлов."],
                    recommendation="Отклонить: файл нечитаемый.",
                    raw_excerpt="",
                )
            else:
                try:
                    raw = _llm_review(text, competition, grade_level, model, api_key, base_url)
                    score = WorkScore(
                        name=name,
                        overall=int(raw.get("overall", 0)),
                        by_criterion={
                            k: int(raw.get("by_criterion", {}).get(k, 0))
                            for k, _ in CHECKLIST
                        },
                        strengths=list(raw.get("strengths", [])),
                        weaknesses=list(raw.get("weaknesses", [])),
                        recommendation=str(raw.get("recommendation", "")),
                        raw_excerpt=text[:300],
                    )
                except Exception as e:  # noqa: BLE001
                    agent.log("error", f"{name}: LLM ошибка: {e}")
                    score = WorkScore(
                        name=name, overall=0,
                        by_criterion={k: 0 for k, _ in CHECKLIST},
                        strengths=[],
                        weaknesses=[f"Ошибка LLM-оценки: {e}"],
                        recommendation="Перепроверить вручную.",
                        raw_excerpt=text[:300],
                    )
            scores.append(score)
            zf.writestr(f"{name}.docx", _per_work_doc(score, competition))
            agent.item_done(
                name,
                summary=f"оценка {score.overall} из 100",
                data={"overall": score.overall, "by_criterion": score.by_criterion},
            )

    out_dir = agent.output_dir
    (out_dir / "per_work.zip").write_bytes(zip_buf.getvalue())

    # Сводный отчёт
    report = Document()
    report.add_heading(f"{competition} — сводное заключение", level=0)
    report.add_paragraph(f"Класс участников: {grade_level}")
    report.add_paragraph(f"Работ принято к проверке: {n}")
    report.add_paragraph(f"Модель эксперта: {model}")

    report.add_heading("Чек-лист научной экспертизы", level=1)
    for _, label in CHECKLIST:
        report.add_paragraph(label, style="List Bullet")

    report.add_heading("Результаты", level=1)
    table = report.add_table(rows=1, cols=2 + len(CHECKLIST))
    hdr = table.rows[0].cells
    hdr[0].text = "Работа"
    for i, (_, label) in enumerate(CHECKLIST):
        hdr[1 + i].text = label
    hdr[-1].text = "Итог"
    for s in sorted(scores, key=lambda x: -x.overall):
        row = table.add_row().cells
        row[0].text = s.name
        for i, (k, _) in enumerate(CHECKLIST):
            row[1 + i].text = str(s.by_criterion.get(k, "—"))
        row[-1].text = str(s.overall)

    report.add_heading("Топ-3 работы", level=1)
    for s in sorted(scores, key=lambda x: -x.overall)[:3]:
        report.add_heading(f"{s.name} — {s.overall}/100", level=2)
        report.add_paragraph(s.recommendation or "—")

    report.save(out_dir / "report.docx")

    agent.progress(1.0, "Готово")
    agent.result(artifacts=[
        {"id": "report", "path": "report.docx"},
        {"id": "per_work", "path": "per_work.zip"},
    ])


if __name__ == "__main__":
    main()
